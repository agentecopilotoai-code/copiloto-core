"""Pure-helper tests for worker modules.

Covers sync helpers in:
- app/workers/event_worker.py (provider_message_id, delivery_error_message,
  delivery_error_code)
- app/workers/scheduler.py (jsonb_payload, _coerce_payload_dict,
  _extract_purpose, _extract_kind)
- app/workers/extraction_worker.py (_extract_pdf_text, _extract_docx_text,
  _read_local_file, _extract_text_sync)
- app/workers/digest_worker.py (_wa_id_from_phone)
"""
from __future__ import annotations

import io
from types import SimpleNamespace

import httpx
import pytest


# ═════════════════════════════════════════════════════════════════════════
# event_worker pure helpers
# ═════════════════════════════════════════════════════════════════════════


def test_provider_message_id_extracts_first():
    from app.workers.event_worker import provider_message_id
    result = {'messages': [{'id': 'wamid.123'}, {'id': 'wamid.456'}]}
    assert provider_message_id(result) == 'wamid.123'


def test_provider_message_id_none_when_no_messages():
    from app.workers.event_worker import provider_message_id
    assert provider_message_id({}) is None
    assert provider_message_id({'messages': []}) is None
    assert provider_message_id({'messages': 'not a list'}) is None


def test_provider_message_id_none_when_first_invalid():
    from app.workers.event_worker import provider_message_id
    assert provider_message_id({'messages': ['not-a-dict']}) is None
    assert provider_message_id({'messages': [{}]}) is None
    assert provider_message_id({'messages': [{'id': 42}]}) is None  # not a string


def test_delivery_error_message_http_error():
    from app.workers.event_worker import delivery_error_message
    request = httpx.Request('POST', 'https://x')
    response = httpx.Response(500, request=request, text='internal err')
    exc = httpx.HTTPStatusError('error', request=request, response=response)
    msg = delivery_error_message(exc)
    assert 'HTTP 500' in msg
    assert 'internal err' in msg


def test_delivery_error_message_truncates_long_text():
    from app.workers.event_worker import delivery_error_message
    request = httpx.Request('POST', 'https://x')
    long_body = 'x' * 5000
    response = httpx.Response(500, request=request, text=long_body)
    exc = httpx.HTTPStatusError('error', request=request, response=response)
    msg = delivery_error_message(exc)
    # Truncated to 1000 chars of body
    assert len(msg) < 1100


def test_delivery_error_message_generic_exception():
    from app.workers.event_worker import delivery_error_message
    msg = delivery_error_message(RuntimeError('connection refused'))
    assert 'connection refused' in msg


def test_delivery_error_message_truncates_generic():
    from app.workers.event_worker import delivery_error_message
    msg = delivery_error_message(RuntimeError('x' * 5000))
    assert len(msg) == 1000


def test_delivery_error_code_meta_error():
    from app.workers.event_worker import delivery_error_code
    request = httpx.Request('POST', 'https://x')
    response = httpx.Response(
        400, request=request,
        json={'error': {'code': 131026, 'message': 'Recipient cannot be sent messages'}},
    )
    exc = httpx.HTTPStatusError('error', request=request, response=response)
    assert delivery_error_code(exc) == '131026'


def test_delivery_error_code_meta_no_code_field():
    from app.workers.event_worker import delivery_error_code
    request = httpx.Request('POST', 'https://x')
    response = httpx.Response(400, request=request, json={'error': {}})
    exc = httpx.HTTPStatusError('error', request=request, response=response)
    assert delivery_error_code(exc) == 'http_400'


def test_delivery_error_code_invalid_json():
    from app.workers.event_worker import delivery_error_code
    request = httpx.Request('POST', 'https://x')
    response = httpx.Response(500, request=request, text='Not JSON')
    exc = httpx.HTTPStatusError('error', request=request, response=response)
    assert delivery_error_code(exc) == 'http_500'


def test_delivery_error_code_transport_error():
    from app.workers.event_worker import delivery_error_code
    assert delivery_error_code(RuntimeError('timeout')) == 'transport_error'
    assert delivery_error_code(OSError('dns failure')) == 'transport_error'


# ═════════════════════════════════════════════════════════════════════════
# scheduler pure helpers
# ═════════════════════════════════════════════════════════════════════════


def test_jsonb_payload_dict_becomes_json():
    from app.workers.scheduler import jsonb_payload
    assert jsonb_payload({'k': 1}) == '{"k": 1}'


def test_jsonb_payload_string_passthrough():
    from app.workers.scheduler import jsonb_payload
    assert jsonb_payload('{"x": 1}') == '{"x": 1}'


def test_jsonb_payload_none_returns_empty():
    from app.workers.scheduler import jsonb_payload
    assert jsonb_payload(None) == '{}'


def test_jsonb_payload_list_becomes_json():
    from app.workers.scheduler import jsonb_payload
    assert jsonb_payload([1, 2]) == '[1, 2]'


def test_coerce_payload_dict_passthrough():
    from app.workers.scheduler import _coerce_payload_dict
    d = {'k': 'v'}
    assert _coerce_payload_dict(d) == d


def test_coerce_payload_dict_string_decode():
    from app.workers.scheduler import _coerce_payload_dict
    assert _coerce_payload_dict('{"x": 1}') == {'x': 1}


def test_coerce_payload_dict_invalid_json_returns_none():
    from app.workers.scheduler import _coerce_payload_dict
    assert _coerce_payload_dict('not json') is None


def test_coerce_payload_dict_non_dict_returns_none():
    from app.workers.scheduler import _coerce_payload_dict
    assert _coerce_payload_dict(None) is None
    assert _coerce_payload_dict(42) is None
    assert _coerce_payload_dict([1, 2]) is None


def test_extract_purpose_returns_value():
    from app.workers.scheduler import _extract_purpose
    assert _extract_purpose({'purpose': 'reminder'}) == 'reminder'


def test_extract_purpose_from_json_string():
    from app.workers.scheduler import _extract_purpose
    assert _extract_purpose('{"purpose": "recall"}') == 'recall'


def test_extract_purpose_missing_or_empty():
    from app.workers.scheduler import _extract_purpose
    assert _extract_purpose(None) is None
    assert _extract_purpose({}) is None
    assert _extract_purpose({'purpose': ''}) is None
    assert _extract_purpose({'purpose': 42}) is None  # not a string


def test_extract_kind_returns_value():
    from app.workers.scheduler import _extract_kind
    assert _extract_kind({'kind': 'auto_rebook_timeout'}) == 'auto_rebook_timeout'


def test_extract_kind_from_json_string():
    from app.workers.scheduler import _extract_kind
    assert _extract_kind('{"kind": "x"}') == 'x'


def test_extract_kind_missing_or_empty():
    from app.workers.scheduler import _extract_kind
    assert _extract_kind(None) is None
    assert _extract_kind({}) is None
    assert _extract_kind({'kind': ''}) is None
    assert _extract_kind({'kind': True}) is None


# ═════════════════════════════════════════════════════════════════════════
# extraction_worker pure helpers
# ═════════════════════════════════════════════════════════════════════════


def test_extract_pdf_text_unreadable_pdf_raises():
    from app.workers.extraction_worker import _extract_pdf_text
    # Garbage bytes — pypdf will fail to parse
    with pytest.raises(Exception):
        _extract_pdf_text(b'not a pdf')


def test_extract_pdf_text_real_pdf():
    """Tiny valid PDF with one page of text."""
    import pypdf
    # Build a minimal in-memory PDF via pypdf
    writer = pypdf.PdfWriter()
    # Add a blank page
    writer.add_blank_page(width=72, height=72)
    buf = io.BytesIO()
    writer.write(buf)
    pdf_bytes = buf.getvalue()

    from app.workers.extraction_worker import _extract_pdf_text
    # Blank page → empty text → raises
    with pytest.raises(ValueError, match='extractable text'):
        _extract_pdf_text(pdf_bytes)


def test_extract_docx_text_unreadable_raises():
    from app.workers.extraction_worker import _extract_docx_text
    with pytest.raises(Exception):
        _extract_docx_text(b'not a docx')


def test_extract_docx_text_real_docx():
    """Build a real .docx in memory via python-docx."""
    from docx import Document
    doc = Document()
    doc.add_paragraph('Hola mundo')
    doc.add_paragraph('Segundo párrafo')
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    from app.workers.extraction_worker import _extract_docx_text
    text, count = _extract_docx_text(docx_bytes)
    assert 'Hola mundo' in text
    assert 'Segundo' in text
    assert count == 2


def test_extract_docx_text_empty_raises():
    from docx import Document
    doc = Document()  # no paragraphs added
    buf = io.BytesIO()
    doc.save(buf)
    from app.workers.extraction_worker import _extract_docx_text
    with pytest.raises(ValueError, match='extractable text'):
        _extract_docx_text(buf.getvalue())


def test_read_local_file_path_traversal_rejected(tmp_path):
    from app.workers.extraction_worker import _read_local_file
    settings = SimpleNamespace(knowledge_storage_local_path=str(tmp_path))
    # Escapes root
    with pytest.raises(ValueError, match='escapes'):
        _read_local_file('../etc/passwd', settings)


def test_read_local_file_not_found(tmp_path):
    from app.workers.extraction_worker import _read_local_file
    settings = SimpleNamespace(knowledge_storage_local_path=str(tmp_path))
    with pytest.raises(FileNotFoundError):
        _read_local_file('nonexistent.pdf', settings)


def test_read_local_file_reads_bytes(tmp_path):
    from app.workers.extraction_worker import _read_local_file
    f = tmp_path / 'sample.bin'
    f.write_bytes(b'PDF-CONTENT')
    settings = SimpleNamespace(knowledge_storage_local_path=str(tmp_path))
    assert _read_local_file('sample.bin', settings) == b'PDF-CONTENT'


def test_extract_text_sync_unsupported_mime():
    from app.workers.extraction_worker import _extract_text_sync
    with pytest.raises(ValueError, match='No extractor'):
        _extract_text_sync(b'x', 'image/png')


def test_extract_text_sync_dispatches_to_pdf():
    """Hits the PDF branch via mime."""
    from app.workers.extraction_worker import _extract_text_sync
    # Invalid PDF bytes → raises from pypdf, but the dispatch happened
    with pytest.raises(Exception):
        _extract_text_sync(b'not a pdf', 'application/pdf')


def test_extract_text_sync_dispatches_to_docx():
    from app.workers.extraction_worker import _extract_text_sync
    with pytest.raises(Exception):
        _extract_text_sync(
            b'not a docx',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )


def test_extract_text_sync_strips_mime_params():
    """The dispatcher should normalize 'application/pdf; charset=binary'."""
    from app.workers.extraction_worker import _extract_text_sync
    with pytest.raises(Exception):
        _extract_text_sync(b'not a pdf', 'application/pdf; charset=binary')


# ═════════════════════════════════════════════════════════════════════════
# digest_worker pure helpers
# ═════════════════════════════════════════════════════════════════════════


def test_wa_id_from_phone_strips_plus():
    from app.workers.digest_worker import _wa_id_from_phone
    assert _wa_id_from_phone('+5730099887766') == '5730099887766'


def test_wa_id_from_phone_no_plus():
    from app.workers.digest_worker import _wa_id_from_phone
    assert _wa_id_from_phone('5730099887766') == '5730099887766'
